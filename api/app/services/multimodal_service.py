"""
多模态文件处理服务

处理图片、文档等多模态文件，转换为 LLM 可用的格式

支持的 Provider:
- DashScope (通义千问): 支持 URL 格式
- Bedrock/Anthropic: 仅支持 base64 格式
- OpenAI: 支持 URL 和 base64 格式
"""
import base64
import csv
import io
import json
import re
import olefile
import struct
import zipfile
from abc import ABC, abstractmethod
from typing import List, Dict, Any, Optional

import PyPDF2
import chardet
import httpx
import magic
import openpyxl
import uuid
from docx import Document
from sqlalchemy.orm import Session

from app.core.config import settings
from app.core.error_codes import BizCode
from app.core.exceptions import BusinessException
from app.core.logging_config import get_business_logger
from app.models import ModelApiKey
from app.models.file_metadata_model import FileMetadata
from app.schemas.app_schema import FileInput, FileType, TransferMethod
from app.schemas.model_schema import ModelInfo
from app.services.audio_transcription_service import AudioTranscriptionService

logger = get_business_logger()

TEXT_MIME = ['text/plain', 'text/x-markdown']
PDF_MIME = ['application/pdf']
DOC_MIME = [
    'application/msword',
    'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
]
XLSX_MIME = [
    'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
    'application/vnd.ms-excel',
]
CSV_MIME = ['text/csv', 'application/csv']
JSON_MIME = ['application/json']


class MultimodalFormatStrategy(ABC):
    """多模态格式策略基类"""

    def __init__(self, file: FileInput):
        self.file = file

    @abstractmethod
    async def format_image(self, url: str, content: bytes | None = None) -> tuple[bool, Dict[str, Any]]:
        """格式化图片"""
        pass

    @abstractmethod
    async def format_document(self, file_name: str, text: str) -> tuple[bool, Dict[str, Any]]:
        """格式化文档"""
        pass

    @abstractmethod
    async def format_audio(self, file_type: str, url: str, content: bytes | None = None) -> tuple[bool, Dict[str, Any]]:
        """格式化音频"""
        pass

    @abstractmethod
    async def format_video(self, url: str) -> tuple[bool, Dict[str, Any]]:
        """格式化视频"""
        pass


class DashScopeFormatStrategy(MultimodalFormatStrategy):
    """通义千问策略"""

    async def format_image(self, url: str, content: bytes | None = None) -> tuple[bool, Dict[str, Any]]:
        """通义千问图片格式：{"type": "image", "image": "url"}"""
        return True, {
            "type": "image",
            "image": url
        }

    async def format_document(self, file_name: str, text: str) -> tuple[bool, Dict[str, Any]]:
        """通义千问文档格式"""
        return True, {
            "type": "text",
            "text": f"<document name=\"{file_name}\">\n文档内容：\n{text}\n</document>"
        }

    async def format_audio(
            self,
            file_type: str,
            url: str,
            content: bytes | None = None,
            transcription: Optional[str] = None
    ) -> tuple[bool, Dict[str, Any]]:
        """
        通义千问音频格式
        - 原生支持: qwen-audio 系列
        - 其他模型: 需要转录为文本
        """
        if transcription:
            return True, {
                "type": "text",
                "text": f"<audio url=\"{url}\">\ntext_transcription:{transcription}\n</audio>"
            }
        # 通义千问音频格式：{"type": "audio", "audio": "url"}
        return True, {
            "type": "audio",
            "audio": url
        }

    async def format_video(self, url: str) -> tuple[bool, Dict[str, Any]]:
        """通义千问视频格式（qwen-vl 系列原生支持）"""
        return True, {
            "type": "video",
            "video": url
        }


class BedrockFormatStrategy(MultimodalFormatStrategy):
    """Bedrock/Anthropic 策略"""

    async def format_image(self, url: str, content: bytes | None = None) -> tuple[bool, Dict[str, Any]]:
        """
        Bedrock/Anthropic 格式: base64 编码
        {"type": "image", "source": {"type": "base64", "media_type": "...", "data": "..."}}
        """

        logger.info(f"下载并编码图片: {url}")

        # 下载图片
        if content is None:
            async with httpx.AsyncClient(timeout=30.0) as client:
                response = await client.get(url, follow_redirects=True)
                response.raise_for_status()
                content = response.content
                self.file.set_content(content)

            # 确定 media type
        content_type = magic.from_buffer(content, mime=True)
        media_type = content_type if content_type.startswith("image/") else "image/jpeg"
        base64_data = base64.b64encode(content).decode("utf-8")

        logger.info(f"图片编码完成: media_type={media_type}, size={len(base64_data)}")

        return True, {
            "type": "image",
            "source": {
                "type": "base64",
                "media_type": media_type,
                "data": base64_data
            }
        }

    async def format_document(self, file_name: str, text: str) -> tuple[bool, Dict[str, Any]]:
        """Bedrock/Anthropic 文档格式（需要 base64 编码）"""
        # Bedrock 文档需要 base64 编码
        text = f"文档内容：\n{text}\n"
        text_bytes = text.encode('utf-8')
        base64_text = base64.b64encode(text_bytes).decode('utf-8')

        return True, {
            "type": "document",
            "source": {
                "type": "base64",
                "media_type": "text/plain",
                "data": base64_text
            }
        }

    async def format_audio(
            self, file_type: str,
            url: str,
            content: bytes | None = None,
            transcription: Optional[str] = None
    ) -> tuple[bool, Dict[str, Any]]:
        """
        Bedrock/Anthropic 音频格式
        不支持原生音频，必须转录为文本
        """
        if transcription:
            return True, {
                "type": "text",
                "text": f"[音频转录]\n{transcription}"
            }
        return False, {
            "type": "text",
            "text": "[音频文件：Bedrock 不支持原生音频，请启用音频转文本功能]"
        }

    async def format_video(self, url: str) -> tuple[bool, Dict[str, Any]]:
        """Bedrock/Anthropic 视频格式"""
        return False, {
            "type": "text",
            "text": f"<video url=\"{url}\">\n[视频文件，当前 provider 暂不支持]\n</video>"
        }


class OpenAIFormatStrategy(MultimodalFormatStrategy):
    """OpenAI 策略"""

    async def format_image(self, url: str, content: bytes | None = None) -> tuple[bool, Dict[str, Any]]:
        """OpenAI 格式: {"type": "image_url", "image_url": {"url": "..."}}"""
        return True, {
            "type": "image_url",
            "image_url": {
                "url": url
            }
        }

    async def format_document(self, file_name: str, text: str) -> tuple[bool, Dict[str, Any]]:
        """OpenAI 文档格式"""
        return True, {
            "type": "text",
            "text": f"<document name=\"{file_name}\">\n文档内容：\n{text}\n</document>"
        }

    async def format_audio(
            self,
            file_type: str,
            url: str,
            content: bytes | None = None,
            transcription: Optional[str] = None
    ) -> tuple[bool, Dict[str, Any]]:
        """
        OpenAI 音频格式
        - gpt-4o-audio 系列支持原生音频（需要 base64 编码）
        - 其他模型使用转录文本
        """
        if transcription:
            return True, {
                "type": "text",
                "text": f"<audio url=\"{url}\">\n{transcription}\n</audio>"
            }

        # OpenAI 音频需要 base64 编码
        try:
            audio_data = content
            if content is None:
                async with httpx.AsyncClient(timeout=30.0) as client:
                    response = await client.get(url, follow_redirects=True)
                    response.raise_for_status()
                    audio_data = response.content
                    self.file.set_content(audio_data)
            base64_audio = base64.b64encode(audio_data).decode('utf-8')

            # 1. 优先从 file_type (MIME) 取扩展名
            file_ext = file_type.split('/')[-1] if file_type and '/' in file_type else None
            # 2. 从响应头 content-type 取
            if not file_ext:
                content_type = magic.from_buffer(audio_data, mime=True)
                file_ext = content_type.split('/')[-1].split(';')[0].strip() if '/' in content_type else None
            # 3. 从 URL 路径取扩展名
            if not file_ext:
                file_ext = url.split('?')[0].rsplit('.', 1)[-1].lower() or None
            # 4. 默认 wav
            # supported_ext = {"wav", "mp3", "mp4", "ogg", "flac", "webm", "m4a", "wave", "x-m4a"}
            file_ext = "wav" if not file_ext else file_ext

            return True, {
                "type": "input_audio",
                "input_audio": {
                    "data": f"data:;base64,{base64_audio}",
                    "format": file_ext
                }
            }
        except Exception as e:
            logger.error(f"下载音频失败: {e}")
            return False, {
                "type": "text",
                "text": f"[音频处理失败: {str(e)}]"
            }

    async def format_video(self, url: str) -> tuple[bool, Dict[str, Any]]:
        """OpenAI 视频格式"""
        return True, {
            "type": "video_url",
            "video_url": {
                "url": url
            }
        }


# Provider 到策略的映射
PROVIDER_STRATEGIES = {
    "dashscope": DashScopeFormatStrategy,
    "bedrock": BedrockFormatStrategy,
    "anthropic": BedrockFormatStrategy,
    "openai": OpenAIFormatStrategy,
    "volcano": OpenAIFormatStrategy,
}


class MultimodalService:
    """
    Service for handling multimodal file processing.

    Attributes:
        db (Session): Database session.
        model_api_key (str): API key for the model provider.
        provider (str): Name of the model provider.
        is_omni (bool): Indicates whether the model supports full multimodal capability.
        capability (list): Capability configuration of the model.
        audio_api_key (str | None): API key used for audio transcription.
        enable_audio_transcription (bool): Whether audio transcription is enabled.
    """

    def __init__(
            self,
            db: Session,
            api_config: ModelInfo | None = None,
            audio_api_key: Optional[str] = None,
            enable_audio_transcription: bool = False,
    ):
        """
        Initialize the multimodal service.

        Args:
            db (Session): Database session.
            api_config (ModelApiKey | None): Model API configuration.
            audio_api_key (str | None): API key for audio transcription.
            enable_audio_transcription (bool): Enable audio transcription.
        """
        self.db = db
        self.api_config = api_config
        if self.api_config is not None:
            self.model_api_key = api_config.api_key
            self.provider = api_config.provider.lower()
            self.is_omni = api_config.is_omni
            self.capability = api_config.capability
        self.audio_api_key = audio_api_key
        self.enable_audio_transcription = enable_audio_transcription

    async def process_files(
            self,
            files: Optional[List[FileInput]],
            workspace_id: uuid.UUID = None,
            document_image_recognition: bool = False,
    ) -> List[Dict[str, Any]]:
        """
        处理文件列表，返回 LLM 可用的格式
        
        Args:
            files: 文件输入列表
            
        Returns:
            List[Dict]: LLM 可用的内容格式列表（根据 provider 返回不同格式）
        """
        if not files:
            return []

        # 获取对应的策略
        # dashscope 的 omni 模型使用 OpenAI 兼容格式
        if self.provider == "dashscope" and self.is_omni:
            strategy_class = OpenAIFormatStrategy
        else:
            strategy_class = PROVIDER_STRATEGIES.get(self.provider)
            if not strategy_class:
                logger.warning(f"未找到 provider '{self.provider}' 的策略，使用默认策略")
                strategy_class = DashScopeFormatStrategy

        result = []
        for idx, file in enumerate(files):
            strategy = strategy_class(file)
            if not file.url:
                file.url = await self.get_file_url(file)
            try:
                if file.type == FileType.IMAGE and "vision" in self.capability:
                    is_support, content = await self._process_image(file, strategy)
                    result.append(content)
                elif file.type == FileType.DOCUMENT:
                    is_support, content = await self._process_document(file, strategy)
                    result.append(content)
                    # 仅当开关开启且模型支持视觉时，才提取文档内嵌图片
                    if document_image_recognition and "vision" in self.capability:
                        img_infos = await self.extract_document_images(file)
                        from app.models.workspace_model import Workspace as WorkspaceModel
                        ws = self.db.query(WorkspaceModel).filter(WorkspaceModel.id == workspace_id).first()
                        tenant_id = ws.tenant_id if ws else None
                        img_result = []
                        for img_info in img_infos:
                            page = img_info["page"]
                            index = img_info["index"]
                            ext = img_info.get("ext", "png")
                            try:
                                _, img_url = await self._save_doc_image_to_storage(img_info["bytes"], ext, tenant_id, workspace_id)
                                placeholder = f"第{page}页 第{index + 1}张" if page > 0 else f"第{index + 1}张"
                                # 在文本内容中追加图片位置标记
                                if result and result[-1].get("type") in ("text", "document"):
                                    key = "text" if "text" in result[-1] else list(result[-1].keys())[-1]
                                    result[-1][key] = result[-1].get(key, "") + f"\n[图片 {placeholder}]: {img_url}"
                                # 将图片以视觉格式追加到消息内容中
                                img_file = FileInput(
                                    type=FileType.IMAGE,
                                    transfer_method=TransferMethod.REMOTE_URL,
                                    url=img_url,
                                    file_type="image/png",
                                )
                                _, img_content = await self._process_image(img_file, strategy_class(img_file))
                                img_result.append(img_content)
                            except Exception as img_err:
                                logger.warning(f"文档图片处理失败: {img_err}")
                        result.extend(img_result)
                elif file.type == FileType.AUDIO and "audio" in self.capability:
                    is_support, content = await self._process_audio(file, strategy)
                    result.append(content)
                elif file.type == FileType.VIDEO and "video" in self.capability:
                    is_support, content = await self._process_video(file, strategy)
                    result.append(content)
                else:
                    logger.warning(f"不支持的文件类型: {file.type}")
            except Exception as e:
                logger.error(
                    f"处理文件失败",
                    extra={
                        "file_index": idx,
                        "file_type": file.type,
                        "error": str(e)
                    },
                    exc_info=True
                )
                # 继续处理其他文件，不中断整个流程
                result.append({
                    "type": "text",
                    "text": f"[文件处理失败: {str(e)}]"
                })

        logger.info(f"成功处理 {len(result)}/{len(files)} 个文件，provider={self.provider}")
        return result

    async def _process_image(self, file: FileInput, strategy) -> tuple[bool, Dict[str, Any]]:
        """
        处理图片文件
        
        Args:
            file: 图片文件输入
            strategy: 格式化策略
            
        Returns:
            Dict: 根据 provider 返回不同格式的图片内容
        """
        try:
            # url = await self.get_file_url(file)
            return await strategy.format_image(file.url, content=file.get_content())
        except Exception as e:
            logger.error(f"处理图片失败: {e}", exc_info=True)
            return False, {
                "type": "text",
                "text": f"[图片处理失败: {str(e)}]"
            }

    async def _process_document(self, file: FileInput, strategy) -> tuple[bool, Dict[str, Any]]:
        """
        处理文档文件（PDF、Word 等）
        
        Returns:
            仅返回文本内容（图片通过 process_files 中的额外步骤追加）
        """
        if file.transfer_method == TransferMethod.REMOTE_URL:
            return True, {
                "type": "text",
                "text": f"<document url=\"{file.url}\">\n{await self.extract_document_text(file)}\n</document>"
            }
        else:
            server_url = settings.FILE_LOCAL_SERVER_URL
            file.url = f"{server_url}/storage/permanent/{file.upload_file_id}"
            text = await self.extract_document_text(file)
            file_metadata = self.db.query(FileMetadata).filter(
                FileMetadata.id == file.upload_file_id
            ).first()
            file_name = file_metadata.file_name if file_metadata else "unknown"
            return await strategy.format_document(file_name, text)

    @staticmethod
    async def _save_doc_image_to_storage(
            img_bytes: bytes,
            ext: str,
            tenant_id: uuid.UUID,
            workspace_id: uuid.UUID,
    ) -> tuple[str, str]:
        """
        将文档内嵌图片保存到存储后端，写入 FileMetadata。

        Returns:
            (file_id_str, permanent_url)
        """
        from app.services.file_storage_service import FileStorageService, generate_file_key
        from app.db import get_db_context

        file_id = uuid.uuid4()
        file_ext = f".{ext}" if not ext.startswith(".") else ext
        content_type = f"image/{ext}"

        file_key = generate_file_key(tenant_id, workspace_id, file_id, file_ext)
        storage_svc = FileStorageService()
        await storage_svc.storage.upload(file_key, img_bytes, content_type)

        with get_db_context() as db:
            meta = FileMetadata(
                id=file_id,
                tenant_id=tenant_id,
                workspace_id=workspace_id,
                file_key=file_key,
                file_name=f"doc_image_{file_id}{file_ext}",
                file_ext=file_ext,
                file_size=len(img_bytes),
                content_type=content_type,
                status="completed",
            )
            db.add(meta)
            db.commit()

        url = f"{settings.FILE_LOCAL_SERVER_URL}/storage/permanent/{file_id}"
        return str(file_id), url

    async def _process_audio(self, file: FileInput, strategy) -> tuple[bool, Dict[str, Any]]:
        """
        处理音频文件
        
        Args:
            file: 音频文件输入
            strategy: 格式化策略
            
        Returns:
            Dict: 根据 provider 返回不同格式的音频内容
        """
        try:
            # url = await self.get_file_url(file)

            # 如果启用音频转文本且有 API Key
            transcription = None
            if self.enable_audio_transcription and self.audio_api_key:
                logger.info(f"开始音频转文本: {file.url}")
                if self.provider == "dashscope":
                    transcription = await AudioTranscriptionService.transcribe_dashscope(file.url, self.audio_api_key)
                elif self.provider == "openai":
                    transcription = await AudioTranscriptionService.transcribe_openai(file.url, self.audio_api_key)
                else:
                    logger.warning(f"Provider {self.provider} 不支持音频转文本")

            return await strategy.format_audio(file.file_type, file.url, file.get_content(), transcription)
        except Exception as e:
            logger.error(f"处理音频失败: {e}", exc_info=True)
            return False, {
                "type": "text",
                "text": f"[音频处理失败: {str(e)}]"
            }

    async def _process_video(self, file: FileInput, strategy) -> tuple[bool, Dict[str, Any]]:
        """
        处理视频文件
        
        Args:
            file: 视频文件输入
            strategy: 格式化策略
            
        Returns:
            Dict: 根据 provider 返回不同格式的视频内容
        """
        try:
            # url = await self.get_file_url(file)
            return await strategy.format_video(file.url)
        except Exception as e:
            logger.error(f"处理视频失败: {e}", exc_info=True)
            return False, {
                "type": "text",
                "text": f"[视频处理失败: {str(e)}]"
            }

    async def get_file_url(self, file: FileInput) -> str:
        """
        获取文件的访问 URL
        
        Args:
            file: File Input Struct
            
        Returns:
            str: 文件访问 URL（永久URL）
            
        Raises:
            BusinessException: 文件不存在
        """
        if file.transfer_method == TransferMethod.REMOTE_URL:
            return file.url
        else:
            file_id = file.upload_file_id

            # 查询 FileMetadata
            file_metadata = self.db.query(FileMetadata).filter(
                FileMetadata.id == file_id,
                FileMetadata.status == "completed"
            ).first()

            if not file_metadata:
                raise BusinessException(
                    f"文件不存在或已删除: {file_id}",
                    BizCode.NOT_FOUND
                )

            # 返回永久URL
            server_url = settings.FILE_LOCAL_SERVER_URL
            return f"{server_url}/storage/permanent/{file_id}"

    async def extract_document_text(self, file: FileInput) -> str:
        """
        提取文档文本内容
        
        Args:
            file: 文件输入
            
        Returns:
            str: 提取的文本内容
        """
        try:
            file_content = file.get_content()
            if not file_content:
                async with httpx.AsyncClient(timeout=30.0) as client:
                    response = await client.get(file.url, follow_redirects=True)
                    response.raise_for_status()
                    file_content = response.content
                    file.set_content(file_content)
            file_mime_type = magic.from_buffer(file_content, mime=True)
            if file_mime_type in TEXT_MIME:
                return self._decode_text_safe(file_content)
            elif file_mime_type in PDF_MIME:
                return await self._extract_pdf_text(file_content)
            elif self._is_word_file(file_content, file_mime_type):
                return await self._extract_word_text(file_content)
            elif self._is_excel_file(file_content, file_mime_type):
                return await self._extract_xlsx_text(file_content)
            elif file_mime_type in CSV_MIME:
                return await self._extract_csv_text(file_content)
            elif file_mime_type in JSON_MIME:
                return await self._extract_json_text(file_content)
            else:
                return f"[Unsupported file type: {file_mime_type}]"
        except Exception as e:
            logger.error(f"Failed to load file. - {e}")
            return "[Failed to load file.]"

    async def extract_document_images(self, file: FileInput) -> list[dict]:
        """
        提取文档中的内嵌图片（支持 PDF 和 DOCX），附带位置信息。

        Returns:
            list[dict]: 每项包含:
                - bytes: 图片二进制
                - page: 所在页码（PDF 从 1 开始，DOCX 为 0）
                - index: 该页/文档内的图片序号（从 0 开始）
                - ext: 图片扩展名（如 png、jpeg）
        """
        try:
            file_content = file.get_content()
            if not file_content:
                async with httpx.AsyncClient(timeout=30.0) as client:
                    response = await client.get(file.url, follow_redirects=True)
                    response.raise_for_status()
                    file_content = response.content
                    file.set_content(file_content)

            file_mime_type = magic.from_buffer(file_content, mime=True)
            if file_mime_type in PDF_MIME:
                return self._extract_pdf_images(file_content)
            elif self._is_word_file(file_content, file_mime_type):
                return self._extract_docx_images(file_content)
            return []
        except Exception as e:
            logger.error(f"提取文档图片失败: {e}")
            return []

    @staticmethod
    def _extract_pdf_images(file_content: bytes) -> list[dict]:
        """从 PDF 提取内嵌图片，附带页码和序号"""
        images = []
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(stream=file_content, filetype="pdf")
            for page_num, page in enumerate(doc, start=1):
                for idx, img in enumerate(page.get_images(full=True)):
                    xref = img[0]
                    base_image = doc.extract_image(xref)
                    images.append({
                        "bytes": base_image["image"],
                        "ext": base_image.get("ext", "png"),
                        "page": page_num,
                        "index": idx,
                    })
            doc.close()
        except ImportError:
            logger.warning("PyMuPDF 未安装，无法提取 PDF 图片，请执行: uv add pymupdf")
        except Exception as e:
            logger.error(f"提取 PDF 图片失败: {e}")
        return images

    @staticmethod
    def _extract_docx_images(file_content: bytes) -> list[dict]:
        """从 DOCX 提取内嵌图片，附带序号（DOCX 无页码概念，page 固定为 0）"""
        images = []
        try:
            if file_content[:2] != b'PK':
                return []
            with zipfile.ZipFile(io.BytesIO(file_content)) as zf:
                media_files = sorted(
                    name for name in zf.namelist()
                    if name.startswith("word/media/") and not name.endswith("/")
                )
                for idx, name in enumerate(media_files):
                    ext = name.rsplit(".", 1)[-1].lower() if "." in name else "png"
                    images.append({
                        "bytes": zf.read(name),
                        "ext": ext,
                        "page": 0,
                        "index": idx,
                    })
        except Exception as e:
            logger.error(f"提取 DOCX 图片失败: {e}")
        return images

    @staticmethod
    async def _extract_pdf_text(file_content: bytes) -> str:
        """提取 PDF 文本"""
        try:
            # 使用 BytesIO 读取 PDF
            text_parts = []
            pdf_file = io.BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            for page in pdf_reader.pages:
                text_parts.append(page.extract_text())
            return '\n'.join(text_parts)
        except Exception as e:
            logger.error(f"提取 PDF 文本失败: {e}")
            return f"[PDF 提取失败: {str(e)}]"

    @staticmethod
    async def _extract_word_text(file_content: bytes) -> str:
        """提取 Word 文档文本（支持 .docx 和旧版 .doc）"""
        # 先尝试 docx（ZIP 格式）
        if file_content[:2] == b'PK':
            try:
                word_file = io.BytesIO(file_content)
                doc = Document(word_file)
                text_lines = []
                for p in doc.paragraphs:
                    text = p.text.strip()
                    if text:
                        text_lines.append(text)

                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            text = cell.text.strip()
                            if text:
                                text_lines.append(text)

                full_text = "\n".join(text_lines)
                return full_text.strip() or "[docx 文件无文本内容]"
            except Exception as e:
                logger.error(f"提取 docx 文本失败: {str(e)}", exc_info=True)
                return f"[docx 提取失败: {str(e)}]"

        # 旧版 .doc（OLE2/CFB 格式），按 Word Binary Format 规范解析 piece table
        try:
            ole = olefile.OleFileIO(io.BytesIO(file_content))
            word_stream = ole.openstream('WordDocument').read()

            # FIB offset 0xA bit9 决定使用 0Table 还是 1Table
            fib_flags = struct.unpack_from('<H', word_stream, 0xA)[0]
            table_name = '1Table' if (fib_flags & 0x0200) else '0Table'
            table_stream = ole.openstream(table_name).read()

            # 从 FIB 读取 fcClx/lcbClx 定位 piece table
            fc_clx, lcb_clx = struct.unpack_from("<II", word_stream, 0x1A2)
            clx = table_stream[fc_clx: fc_clx + lcb_clx]

            # 解析 CLX，找到 PlcPcd（piece table）
            i, plc_pcd = 0, None
            while i < len(clx):
                clxt = clx[i]
                if clxt == 0x01:
                    i += 3 + struct.unpack_from('<H', clx, i + 1)[0]
                elif clxt == 0x02:
                    cb = struct.unpack_from('<I', clx, i + 1)[0]
                    plc_pcd = clx[i + 5: i + 5 + cb]
                    break
                else:
                    break

            if plc_pcd is None:
                raise ValueError("PlcPcd not found")

            # PlcPcd: (n+1) 个 CP（4字节）+ n 个 PCD（8字节）
            n_pieces = (len(plc_pcd) - 4) // 12
            cp_array = [struct.unpack_from('<I', plc_pcd, k * 4)[0] for k in range(n_pieces + 1)]

            parts = []
            for k in range(n_pieces):
                fc_value = struct.unpack_from('<I', plc_pcd, (n_pieces + 1) * 4 + k * 8 + 2)[0]
                is_ansi = bool(fc_value & 0x40000000)
                fc = fc_value & 0x3FFFFFFF
                char_count = cp_array[k + 1] - cp_array[k]

                if is_ansi:
                    parts.append(word_stream[fc: fc + char_count].decode('cp1252', errors='replace'))
                else:
                    parts.append(word_stream[fc: fc + char_count * 2].decode('utf-16-le', errors='replace'))

            ole.close()
            result = re.sub(r'[\x00-\x1f\x7f]', '', ''.join(parts))
            return result.strip()

        except Exception as e:
            logger.error(f"提取 doc 文本失败: {e}")
            return f"[doc 提取失败: {str(e)}]"

    @staticmethod
    async def _extract_xlsx_text(file_content: bytes) -> str:
        """提取 Excel 文本（支持 .xlsx 和旧版 .xls）"""
        # xlsx（ZIP 格式）
        if file_content[:2] == b'PK':
            try:
                wb = openpyxl.load_workbook(io.BytesIO(file_content), read_only=True, data_only=True)
                parts = []
                for sheet in wb.worksheets:
                    parts.append(f"[Sheet: {sheet.title}]")
                    for row in sheet.iter_rows(values_only=True):
                        parts.append('\t'.join('' if v is None else str(v) for v in row))
                return '\n'.join(parts)
            except Exception as e:
                logger.error(f"提取 xlsx 文本失败: {e}")
                return f"[xlsx 提取失败: {str(e)}]"

        # xls（OLE2/BIFF 格式）
        try:
            import xlrd
            wb = xlrd.open_workbook(file_contents=file_content)
            parts = []
            for sheet in wb.sheets():
                parts.append(f"[Sheet: {sheet.name}]")
                for row_idx in range(sheet.nrows):
                    parts.append('\t'.join(str(sheet.cell_value(row_idx, col)) for col in range(sheet.ncols)))
            return '\n'.join(parts)
        except Exception as e:
            logger.error(f"提取 xls 文本失败: {e}")
            return f"[xls 提取失败: {str(e)}]"

    async def _extract_csv_text(self, file_content: bytes) -> str:
        """提取 CSV 文本"""
        try:
            text = self._decode_text_safe(file_content)
            reader = csv.reader(io.StringIO(text))
            return '\n'.join('\t'.join(row) for row in reader)
        except Exception as e:
            logger.error(f"提取 CSV 文本失败: {e}")
            return f"[CSV 提取失败: {str(e)}]"

    async def _extract_json_text(self, file_content: bytes) -> str:
        """提取 JSON 文本"""
        try:
            text = self._decode_text_safe(file_content)
            data = json.loads(text)
            return json.dumps(data, ensure_ascii=False, indent=2)
        except Exception as e:
            logger.error(f"提取 JSON 文本失败: {e}")
            return f"[JSON 提取失败: {str(e)}]"

    def _is_word_file(self, file_content: bytes, mime_type: str) -> bool:
        """判断是不是 Word 文件（doc / docx），不依赖后缀"""
        # 旧版 .doc
        if mime_type == 'application/msword':
            return True

        # 新版 .docx（ZIP 内部包含 word/document.xml）
        header = file_content[:4]
        if header == b'PK\x03\x04':
            try:
                with zipfile.ZipFile(io.BytesIO(file_content)) as zf:
                    return "word/document.xml" in zf.namelist()
            except:
                pass

        return False

    def _is_excel_file(self, file_content: bytes, mime_type: str) -> bool:
        """判断是不是 Excel 文件（xls / xlsx），不依赖后缀"""
        # 旧版 .xls
        if mime_type == 'application/vnd.ms-excel':
            return True

        # 新版 .xlsx（ZIP 内部包含 xl/workbook.xml）
        header = file_content[:4]
        if header == b'PK\x03\x04':
            try:
                with zipfile.ZipFile(io.BytesIO(file_content)) as zf:
                    return "xl/workbook.xml" in zf.namelist()
            except:
                pass

        return False

    @staticmethod
    def _decode_text_safe(file_content: bytes) -> str:
        """
        【万能文本解码】
        自动检测编码，支持 utf-8 / gbk / gb2312 / utf-8-sig / ascii 等
        永远不报错，永远不乱码
        """
        if not file_content:
            return ""

        # 1. 自动检测文件编码
        detect = chardet.detect(file_content)
        encoding = detect.get("encoding") or "utf-8"
        encoding = encoding.lower()

        # 2. 兼容常见中文编码
        compatible_encodings = ["utf-8", "gbk", "gb18030", "gb2312", "ascii", "latin-1"]

        # 3. 按优先级尝试解码
        for enc in [encoding] + compatible_encodings:
            if not enc:
                continue
            try:
                return file_content.decode(enc.strip())
            except (UnicodeDecodeError, LookupError):
                continue

        # 终极兜底
        return file_content.decode("utf-8", errors="replace")


def get_multimodal_service(db: Session) -> MultimodalService:
    """获取多模态服务实例（依赖注入）"""
    return MultimodalService(db)
